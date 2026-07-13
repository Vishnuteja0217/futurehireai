from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import OpenAI
from dotenv import load_dotenv
from resume_parser import extract_resume_text
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import inch
import uuid

import os
import json
import requests
from bs4 import BeautifulSoup

load_dotenv()

app = FastAPI()

# Enable frontend access
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Supabase client for caching job search results.
# We use the SERVICE KEY (not anon key) because this is server-side
# and needs to bypass RLS on internal tables like job_search_cache.
from supabase import create_client as create_supabase_client
supabase = create_supabase_client(
    os.getenv("SUPABASE_URL"),
    os.getenv("SUPABASE_SERVICE_KEY"),
)

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')

    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D1D5DB')

    border.append(bottom)
    pPr.append(border)

# Request schema
class JDRequest(BaseModel):
    job_description: str

class ResumeJDRequest(BaseModel):
    resume_text: str
    job_description: str

class TailorResumeRequest(BaseModel):
    """
    Used by /generate-tailored-resume only. Adds an optional list of
    skills the candidate has CONFIRMED they actually used (collected
    via the SkillConfirmationModal on the frontend). These confirmed
    skills get woven into the tailored resume.
    """
    resume_text: str
    job_description: str
    confirmed_skills: list[str] = []

class AnswerEvaluationRequest(BaseModel):
    question: str
    answer: str
    job_description: str

@app.get("/")
def home():
    return {
        "message": "AI Career Copilot Backend Running"
    }


# Resume Upload Endpoint
@app.post("/upload-resume")
async def upload_resume(file: UploadFile = File(...)):

    file_location = f"temp_{file.filename}"

    with open(file_location, "wb") as f:
        f.write(await file.read())

    extracted_text = extract_resume_text(file_location)

    return {
        "resume_text": extracted_text
    }


# JD Analysis Endpoint
@app.post("/analyze-jd")
def analyze_jd(data: JDRequest):

    prompt = f"""
You are an expert recruiter and career coach.

Analyze the following job description.

Return ONLY valid JSON.

Format:

{{
  "required_skills": [],
  "hidden_expectations": [],
  "interview_topics": [],
  "preparation_questions": [],
  "study_topics": [],
  "high_risk_areas": [],
  "priority_topics": []
}}

Rules:
- Each field must contain concise bullet points
- No markdown
- No explanations outside JSON
- No extra text

Job Description:
{data.job_description}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an expert recruiter and career coach."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    parsed_response = json.loads(
        response.choices[0].message.content
    )

    return parsed_response


@app.post("/compare-resume-jd")
def compare_resume_jd(data: ResumeJDRequest):

    prompt = f"""
You are an experienced senior hiring manager interviewing this candidate for the role.

Do NOT summarize the resume.

Your task is to critically evaluate the candidate.

Think like a skeptical technical interviewer.

CRITICAL — required vs preferred qualifications:
Most JDs separate qualifications into REQUIRED (must-have) and PREFERRED 
(nice-to-have / bonus). You MUST distinguish these when scoring:

- REQUIRED skills: heavily weighted (~80% of match_score). Missing required 
  skills is a major penalty.
- PREFERRED skills: lightly weighted (~20%). Missing preferred skills is a 
  small or zero penalty. Having preferred skills is a bonus.

Look for JD section headers like: "Required", "Must have", "Qualifications", 
"Minimum qualifications", "Basic qualifications" → REQUIRED.

And: "Preferred", "Nice to have", "Bonus", "Plus", "Preferred qualifications",
"Desirable" → PREFERRED.

If the JD doesn't separate them, treat the FIRST listed qualifications as 
required and any later "preferred" or "bonus" mentions as preferred.

SCORING RULES with this lens:
- Candidate has ALL required + some preferred → 85-95
- Candidate has ALL required + zero preferred → 80-88
- Candidate has 80% of required → 65-78
- Candidate has 50% of required → 40-60
- Candidate missing many required → 20-50

A candidate who matches all required qualifications should NEVER score 
below 78, regardless of how many preferred items they're missing. 
"Missing preferred" is not a critical gap — it's just not a bonus.

When listing critical_gaps and recruiter_concerns, ONLY include things 
that are genuinely required for the role. Do not list missing preferred 
qualifications as critical gaps.

Analyze:
- what experience appears weak or shallow
- what recruiter concerns exist
- what the resume FAILS to prove
- which JD expectations are insufficiently demonstrated
- Technical depth concerns: identify areas where the candidate's resume looks shallow, incomplete, or vulnerable to deeper technical questioning. Explain what interviewers may probe further and why. Do NOT list generic skills.
- Hidden role expectations: infer unstated expectations recruiters and hiring managers likely have based on the JD, role seniority, tech stack, and industry norms. Focus on ownership, execution, communication, reliability, production thinking, and collaboration.
- what security/compliance gaps are visible
- which resume bullets sound generic or low-impact
- how the candidate could better position existing experience

Avoid generic observations.

Provide deep recruiter insight.

Return ONLY valid JSON.

The JSON response MUST include match_score, recruiter_confidence, and recruiter_verdict at the top level.

Format:

{{
  "match_score": 0,
  "recruiter_confidence": "",
  "recruiter_verdict": "",
  "top_strengths": [],
  "critical_gaps": [],
  "recruiter_concerns": [],
  "likely_rejection_reasons": [],
  "technical_depth_concerns": [],
  "hidden_role_expectations": [],
  "resume_bullet_improvements": [],
  "high_priority_study_areas": [],
  "likely_interviewer_focus_areas": [],
  "scenario_based_questions": [],
  "common_interview_questions": [],
  "mock_interview_questions": []
}}

- match_score must be an integer from 0-100.
- recruiter_confidence must be one of:
  Low
  Medium
  Medium-High
  High
- recruiter_verdict must be a concise recruiter-style evaluation.

Rules:
- Be brutally realistic
- Think deeply
- Avoid repeating resume content
- First identify the target role/domain from the JD, such as Java backend, frontend, data science, DevOps, cloud, cybersecurity, AI/ML, product, business analyst, QA, or other.
- All gaps, concerns, resume improvements, study areas, and scenario questions MUST be specific to that identified role.
- Do NOT generate DevOps, cloud, Kubernetes, security, or infrastructure questions unless the JD actually requires those areas.
- If the JD is for a Java role, focus on Java, Spring Boot, REST APIs, databases, OOP, system design, testing, concurrency, microservices, debugging, and backend engineering expectations.
- If the JD is for another role, adapt completely to that role.
- Avoid irrelevant technologies or interview topics not supported by the JD.
- Avoid generic ATS-style observations
- Focus on what is MISSING or weak
- Resume improvements should sound like strong recruiter-quality bullets
- Do not invent fake experience
- Focus heavily on interview readiness and recruiter perception
- No markdown
- No explanations outside JSON
- initial_ats_score should estimate the current resume-to-JD match from 0 to 100
- initial_ats_reasoning should explain the score using concise bullet points
- Base initial ATS score on keyword alignment, role relevance, experience depth, missing JD requirements, and resume clarity
- Do not inflate the score unrealistically
- For technical_depth_concerns, each item must explain WHY it is a concern, not just name a technology.
- For hidden_role_expectations, each item must describe an inferred expectation that is not directly stated in the JD.
- Avoid generic keyword lists.
- Write like a skeptical hiring manager giving useful private feedback.
- match_score should be a realistic ATS/recruiter alignment score from 0-100.
- recruiter_confidence should be one of: Low, Medium, Medium-High, High.
- recruiter_verdict should sound like a real recruiter summary, balancing strengths and concerns honestly.
- Strong resumes should receive positive reinforcement instead of unnecessary criticism.

Resume:
{data.resume_text}

Job Description:
{data.job_description}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a senior recruiter and interview strategist."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    parsed_response = json.loads(
        response.choices[0].message.content
    )

    if "match_score" not in parsed_response:
        parsed_response["match_score"] = parsed_response.get("initial_ats_score", 0)

    if "recruiter_confidence" not in parsed_response:
        parsed_response["recruiter_confidence"] = "Medium"

    if "recruiter_verdict" not in parsed_response:
        parsed_response["recruiter_verdict"] = "Resume analyzed successfully. Review the detailed insights below."

    return parsed_response

@app.post("/evaluate-answer")
def evaluate_answer(data: AnswerEvaluationRequest):

    prompt = f"""
You are a realistic but supportive interview evaluator.

Your job is to evaluate the candidate's answer honestly.

Important:
Do NOT automatically generate criticism.
Do NOT force a "better answer" if the candidate's answer is already strong.
First decide the actual quality of the answer.

Return ONLY valid JSON.

Format:

{{
  "answer_quality": "",
  "score": 0,
  "is_interview_ready": false,
  "strengths": [],
  "improvements": [],
  "missing_concepts": [],
  "better_answer": "",
  "follow_up_question": ""
}}

Evaluation rules:
- answer_quality must be one of: "weak", "decent", "strong", "excellent"
- score must be between 0 and 100
- If the answer is excellent, improvements can be minor and missing_concepts can be empty
- If the answer is excellent, better_answer should say: "Your answer is already strong. No major rewrite needed."
- If the answer is strong, only suggest small polish improvements
- If the answer is weak or vague, clearly explain what is missing
- Be honest, but motivating
- Do not be harsh
- Do not sugarcoat
- Evaluate based on role expectations from the job description
- Check for technical depth, clarity, structure, real-world thinking, and interview readiness
- Follow-up question should probe deeper only if useful

Job Description:
{data.job_description}

Interview Question:
{data.question}

Candidate Answer:
{data.answer}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a fair, realistic, and supportive interview evaluator."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    parsed_response = json.loads(
        response.choices[0].message.content
    )

    print(parsed_response)

    return parsed_response

@app.post("/detect-missing-skills")
def detect_missing_skills(data: ResumeJDRequest):
    """
    Identifies specific tools/technologies the JD mentions that are NOT
    present in the candidate's resume. Used to ask the candidate
    "have you used these?" before tailoring, so we can honestly
    add real experience the candidate forgot to list.
    
    We deliberately return ONLY tool/technology names — no fluff,
    no explanations, no soft skills. Frontend uses these as checklist items.
    """

    prompt = f"""
You are an ATS analyst comparing a job description to a candidate's resume.

Identify SPECIFIC TOOLS and TECHNOLOGIES that the JD mentions but the
candidate's resume does NOT clearly include. Return ONLY the tools/technologies
that are MISSING from the resume.

STRICT RULES (read carefully — items must pass ALL these tests):

A "tool" qualifies ONLY if it is a SPECIFIC NAMED product/library/service
that a developer would either use or not use. It must be answerable as a 
clear binary: "yes I used X" or "no I haven't."

GOOD examples (return these):
- Terraform (specific tool)
- Kubernetes (specific platform)
- PostgreSQL (specific database)
- GitLab CI (specific platform)
- Jenkins (specific tool)
- AWS Lambda (specific service)
- React, Vue, Angular (specific frameworks)
- Spring Boot (specific framework)
- Docker (specific tool)
- Redis (specific tool)

BAD examples — NEVER return these, even if the JD mentions them:
- Certifications (CCNA, CCNP, AWS Certified, Microsoft Certified, etc.)
- Abstract categories ("firewalls", "logging tools", "alerting tools",
  "monitoring solutions", "databases", "cloud platforms")
- Conceptual capabilities ("hybrid infrastructure", "enterprise networking",
  "high availability", "scalability", "disaster recovery")
- Networking primitives that EVERY engineer has used and can't differentiate
  on: "subnets", "DNS", "VPNs", "load balancers", "firewalls" (generic),
  "TCP/IP", "HTTP", "TLS", "SSL", "routing", "switching"
- Cloud features that are part of using the cloud, not differentiating skills:
  "VNets", "VPCs", "Network Security Groups", "NSGs", "Security Groups",
  "IAM roles", "subnets", "route tables"
- Methodologies (Agile, DevOps, Scrum, Waterfall)
- Generic concepts (CI/CD, microservices, RESTful APIs, OOP, REST, JSON, XML)
- Soft skills (leadership, communication, problem-solving)
- Seniority levels or years of experience
- Role responsibilities
- Generic verbs/actions ("monitoring", "automation", "scripting", "deployment")

If you're unsure whether an item is a "tool" — apply BOTH tests:

TEST 1 — Specificity test:
  Could a developer say "I used <X> for 6 months on a project"? 
  If yes → continue to test 2. If awkward → skip.

TEST 2 — Differentiation test:
  Would saying "I used X" on a resume meaningfully differentiate this
  candidate from any other? Or has EVERY developer used it?
  - "I used Terraform" → differentiating, not everyone has → KEEP
  - "I used DNS" → everyone has → SKIP
  - "I used Subnets" → everyone has → SKIP
  - "I used Bicep" → Azure-specific, differentiating → KEEP
  - "I used Akamai" → specific CDN product → KEEP

Both tests must pass. If either fails, exclude the item.

Additional rules:
- DO NOT return tools already mentioned in the resume.
- Maximum 8 items.
- Items must be SPECIFIC NAMES, not categories.
- If resume mentions a related but different tool (e.g., resume has
  CloudFormation, JD asks Terraform), include the JD's tool as missing.
- If the JD has fewer than 3 specific tool names that pass these rules,
  return an empty list. Better to skip the modal than ask bad questions.

Return ONLY valid JSON. No markdown, no explanations.

Format:

{{
  "missing_skills": []
}}

Resume:
{data.resume_text}

Job Description:
{data.job_description}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an honest ATS analyst comparing resumes to job descriptions."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    parsed_response = json.loads(
        response.choices[0].message.content
    )

    # Defensive: ensure missing_skills is always a list of clean strings,
    # max 10 items. Frontend assumes this shape.
    raw_skills = parsed_response.get("missing_skills", [])
    if not isinstance(raw_skills, list):
        raw_skills = []

    clean_skills = []
    for skill in raw_skills:
        if isinstance(skill, str) and skill.strip():
            clean_skills.append(skill.strip())
        if len(clean_skills) >= 10:
            break

    return {"missing_skills": clean_skills}

# ============================================================================
# Job Search Endpoint (/search-jobs)
# ============================================================================
# Uses JSearch API (RapidAPI) to find jobs and enriches results with H1B
# sponsor data from USCIS FY2025 LCA disclosure data.
#
# Flow:
#   1. Frontend sends { role, location, experience_level }
#   2. We call JSearch to get raw job listings
#   3. For each job, we look up the employer in h1b_sponsors.json
#      to determine H1B sponsorship history
#   4. We sort so H1B sponsors appear FIRST
#   5. Return enriched, sorted jobs to frontend
# ============================================================================

# Load H1B sponsor data ONCE at startup (not on every request).
# This is a ~2MB JSON with 62,983 companies mapped to their H1B filing counts.
_H1B_SPONSORS_PATH = os.path.join(
    os.path.dirname(__file__), "data", "h1b_sponsors.json"
)

try:
    with open(_H1B_SPONSORS_PATH, "r", encoding="utf-8") as f:
        H1B_SPONSORS = json.load(f)
    print(f"Loaded {len(H1B_SPONSORS):,} H1B sponsors")
except FileNotFoundError:
    print(f"WARNING: h1b_sponsors.json not found at {_H1B_SPONSORS_PATH}")
    H1B_SPONSORS = {}


def normalize_company_name(name: str) -> str:
    """Normalize company name for H1B lookup.
    
    JSearch returns names like 'Microsoft', 'Microsoft Corporation',
    'Microsoft Corp' — but our USCIS data uses 'MICROSOFT CORPORATION'.
    We uppercase everything and strip whitespace to improve match rate.
    """
    if not name:
        return ""
    return name.upper().strip()


def check_h1b_sponsor(company_name: str) -> dict:
    """Check if a company sponsors H1B visas.
    
    Returns:
        {
            "is_sponsor": bool,        # True if company found in USCIS data
            "filings_count": int,      # Number of H1B filings in FY2025
        }
    
    Matching strategy:
        1. Exact match on normalized name (fastest, most confident)
        2. Substring match — check if any H1B sponsor contains this name
           (catches "Google" matching "GOOGLE LLC")
    """
    normalized = normalize_company_name(company_name)
    if not normalized:
        return {"is_sponsor": False, "filings_count": 0}
    
    # Strategy 1: exact match
    if normalized in H1B_SPONSORS:
        return {
            "is_sponsor": True,
            "filings_count": H1B_SPONSORS[normalized],
        }
    
    # Strategy 2: substring match (Google → GOOGLE LLC)
    # We check if the JSearch company name appears at the start of any 
    # sponsor name — avoids false matches like "APPLE" matching "PINEAPPLE INC"
    for sponsor_name, count in H1B_SPONSORS.items():
        if sponsor_name.startswith(normalized + " ") or sponsor_name == normalized:
            return {
                "is_sponsor": True,
                "filings_count": count,
            }
    
    return {"is_sponsor": False, "filings_count": 0}

import hashlib
from datetime import datetime, timezone, timedelta


def build_cache_key(role: str, location: str, exp: str) -> str:
    """Build a stable hash key for cache lookups.
    
    Normalizes inputs (lowercase, strip) so 'Software Engineer' and
    'software engineer  ' hit the same cache entry.
    """
    normalized = f"{role.lower().strip()}|{location.lower().strip()}|{exp.lower().strip()}"
    # SHA-256 (not for security — just a stable content-addressable cache key).
    # We use SHA over MD5 to keep security scanners happy; the crypto strength
    # doesn't matter because cache keys aren't secret.
    return hashlib.sha256(normalized.encode()).hexdigest()


# 24-hour cache TTL. Jobs change daily on major boards.
CACHE_TTL_HOURS = 24

class JobSearchRequest(BaseModel):
    role: str                     # e.g. "Software Engineer"
    location: str = ""            # optional; e.g. "Austin, TX" or "United States"
    experience_level: str = ""    # optional; e.g. "entry_level", "mid_level", "senior_level"


@app.post("/search-jobs")
def search_jobs(data: JobSearchRequest):
    """Search for jobs and mark H1B sponsors.
    
    Flow:
      1. Check Supabase cache — if hit within 24h, return cached result
      2. Cache miss → call JSearch API
      3. Enrich with H1B sponsor data
      4. Sort (H1B sponsors first)
      5. Store in cache for next 24h
      6. Return
    
    Cache is keyed on (role + location + exp). Same query from any user
    within 24h reuses the cached result — dramatically extends the 200
    free JSearch calls/month.
    """
    cache_key = build_cache_key(data.role, data.location, data.experience_level)
    
    # ---- STEP 1: Check cache ----
    try:
        cached = supabase.table("job_search_cache") \
            .select("jobs_json, cached_at") \
            .eq("cache_key", cache_key) \
            .maybe_single() \
            .execute()
        
        if cached and cached.data:
            cached_at = datetime.fromisoformat(cached.data["cached_at"].replace("Z", "+00:00"))
            age = datetime.now(timezone.utc) - cached_at
            
            if age < timedelta(hours=CACHE_TTL_HOURS):
                # Still fresh — return cached
                jobs = cached.data["jobs_json"]
                return {
                    "jobs": jobs,
                    "total_returned": len(jobs),
                    "h1b_sponsors_count": sum(1 for j in jobs if j.get("h1b_sponsor")),
                    "query": f"{data.role} in {data.location}".strip(),
                    "cached": True,
                    "cache_age_hours": round(age.total_seconds() / 3600, 1),
                }
    except Exception as e:
        # Cache is best-effort — if it fails, just skip and hit JSearch
        print(f"Cache read failed (non-fatal): {e}")
    
    # ---- STEP 2: Cache miss → call JSearch ----
    api_key = os.getenv("JSEARCH_API_KEY")
    if not api_key:
        return {"error": "JSearch API key not configured", "jobs": []}
    
    query_parts = [data.role]
    if data.location:
        query_parts.append(f"in {data.location}")
    query_string = " ".join(query_parts).strip()
    
    url = "https://jsearch.p.rapidapi.com/search"
    headers = {
        "X-RapidAPI-Key": api_key,
        "X-RapidAPI-Host": "jsearch.p.rapidapi.com",
    }
    params = {
        "query": query_string,
        "page": "1",
        "num_pages": "1",
        "country": "us",
        "date_posted": "month",
    }
    if data.experience_level:
        params["job_requirements"] = data.experience_level
    
    try:
        response = requests.get(url, headers=headers, params=params, timeout=15)
        response.raise_for_status()
        jsearch_data = response.json()
    except requests.exceptions.RequestException as e:
        return {"error": f"JSearch API failed: {str(e)}", "jobs": []}
    
    raw_jobs = jsearch_data.get("data", [])
    
    # ---- STEP 3: Enrich with H1B data ----
    enriched_jobs = []
    for job in raw_jobs:
        employer = job.get("employer_name", "")
        h1b_info = check_h1b_sponsor(employer)
        
        enriched_jobs.append({
            "job_id": job.get("job_id", ""),
            "title": job.get("job_title", ""),
            "company": employer,
            "company_logo": job.get("employer_logo", ""),
            "location": _format_location(job),
            "job_type": job.get("job_employment_type", ""),
            "is_remote": job.get("job_is_remote", False),
            "posted_at": job.get("job_posted_at_datetime_utc", ""),
            "salary_min": job.get("job_min_salary"),
            "salary_max": job.get("job_max_salary"),
            "salary_period": job.get("job_salary_period", ""),
            "description_snippet": _truncate(job.get("job_description", ""), 400),
            "description_full": job.get("job_description", ""),
            "apply_url": job.get("job_apply_link", ""),
            "source": job.get("job_publisher", ""),
            "h1b_sponsor": h1b_info["is_sponsor"],
            "h1b_filings_2025": h1b_info["filings_count"],
        })
    
    # ---- STEP 4: Sort (H1B sponsors first) ----
    enriched_jobs.sort(
        key=lambda j: (
            0 if j["h1b_sponsor"] else 1,
            -j["h1b_filings_2025"],
        )
    )
    
    # ---- STEP 5: Store in cache (best-effort, don't fail request) ----
    try:
        supabase.table("job_search_cache").upsert({
            "cache_key": cache_key,
            "query_role": data.role,
            "query_location": data.location,
            "query_exp": data.experience_level,
            "jobs_json": enriched_jobs,
            "cached_at": datetime.now(timezone.utc).isoformat(),
        }).execute()
    except Exception as e:
        print(f"Cache write failed (non-fatal): {e}")
    
    # ---- STEP 6: Return ----
    return {
        "jobs": enriched_jobs,
        "total_returned": len(enriched_jobs),
        "h1b_sponsors_count": sum(1 for j in enriched_jobs if j["h1b_sponsor"]),
        "query": query_string,
        "cached": False,
    }

@app.get("/job-details/{job_id}")
def get_job_details(job_id: str):
    """
    Fetch a single job by ID from the Supabase cache.
    
    Detail pages call this when a user hits /app/jobs/[job_id] directly
    (new tab, shared URL, refresh) and doesn't have the job in sessionStorage.
    
    Scans all recent cached searches for the job — free but limited to
    jobs that appeared in a search within the last 24h. Anything older
    returns 404. That's acceptable for v1 — most shared job links get 
    clicked quickly.
    """
    try:
        # URL-decode the job_id — frontend encodes it because JSearch IDs
        # contain characters like = and / that break URL routing.
        from urllib.parse import unquote
        decoded_id = unquote(job_id)
        
        # Query all rows in job_search_cache and scan for the job
        cached_rows = supabase.table("job_search_cache") \
            .select("jobs_json") \
            .execute()
        
        if not cached_rows.data:
            return {"error": "Job not found. Try searching again."}
        
        # Scan every cached search for the matching job_id
        for row in cached_rows.data:
            jobs = row.get("jobs_json", [])
            if not jobs:
                continue
            for job in jobs:
                if job.get("job_id") == decoded_id:
                    return {"job": job}
        
        # Nothing found — job is either fake, from a very old search, or the cache expired
        return {"error": "Job not found. Try searching again."}
    
    except Exception as e:
        print(f"/job-details error: {e}")
        return {"error": "Failed to fetch job details."}

def _format_location(job: dict) -> str:
    """Combine city, state, country into a readable location string."""
    parts = [
        job.get("job_city", ""),
        job.get("job_state", ""),
        job.get("job_country", ""),
    ]
    return ", ".join(p for p in parts if p)
 

def _truncate(text: str, max_length: int) -> str:
    """Truncate text with ellipsis if too long."""
    if not text:
        return ""
    if len(text) <= max_length:
        return text
    return text[:max_length].rsplit(" ", 1)[0] + "..."

@app.post("/generate-tailored-resume")
def generate_tailored_resume(data: TailorResumeRequest):

    # ----------------------------------------------------------------
    # STEP 1: Score the ORIGINAL resume first.
    # We do this so the tailoring prompt knows EXACTLY what the
    # scorer flagged as missing. Targeted improvement >> generic improvement.
    # ----------------------------------------------------------------
    original_score_data = compare_resume_jd(data)
    original_score = original_score_data.get("match_score", 0)
    original_gaps = original_score_data.get("critical_gaps", [])
    original_concerns = original_score_data.get("recruiter_concerns", [])
    original_improvements = original_score_data.get("resume_bullet_improvements", [])

    def _format_bullets(items):
        """Convert a list to numbered bullets for the prompt."""
        if not items:
            return "(none flagged)"
        return "\n".join(
            f"{i+1}. {item}" for i, item in enumerate(items) if isinstance(item, str) and item.strip()
        )

    # ----------------------------------------------------------------
    # STEP 2: Build a TARGETED tailoring prompt.
    # Pass the original score + recruiter's specific feedback so the
    # rewrite addresses real weaknesses, not invented ones.
    # ----------------------------------------------------------------
    prompt = f"""
You are an expert ATS analyst, technical recruiter, and resume strategist.

The candidate's CURRENT resume scored {original_score}/100 against this job description.

A senior recruiter reviewed the resume and flagged these specific issues:

CRITICAL GAPS (must address):
{_format_bullets(original_gaps)}

RECRUITER CONCERNS (worth addressing):
{_format_bullets(original_concerns)}

SUGGESTED BULLET IMPROVEMENTS (recruiter's exact recommendations):
{_format_bullets(original_improvements)}

YOUR JOB: Rewrite the resume to specifically address these issues using ONLY
the candidate's real existing experience. The goal is to push the score
ABOVE the current {original_score} by addressing recruiter feedback.

CONFIRMED ADDITIONAL EXPERIENCE:
The candidate has explicitly CONFIRMED hands-on experience with the following
tools/technologies that were not currently on their resume (the system asked
the candidate "have you used these?" and they answered YES):

{_format_bullets(data.confirmed_skills) if data.confirmed_skills else "(none — candidate did not confirm any additional tools)"}

You MAY add these confirmed skills to the resume since the candidate
verified they have actually used them. Weave them naturally into:
1. The Technical Skills section (always add to the appropriate category)
2. Existing bullets where the work described could plausibly involve them
   (e.g., if they confirmed "Terraform" and have a bullet about "AWS
   infrastructure," you may say "Designed AWS infrastructure using Terraform")

Do NOT add the confirmed skills to bullets where they don't fit the work
described. Do NOT invent NEW bullets just to feature the confirmed skills —
weave them into existing experience.

VERY IMPORTANT RULES:
- Do NOT invent fake experience, tools, projects, metrics, or certifications.
- Do NOT add technologies the candidate cannot justify from their original resume.
- Only reframe and emphasize what is already real in the resume.
- Specifically address the gaps and concerns listed above.
- Strengthen bullets using action + technology + measurable impact.
- Add JD-relevant keywords naturally where the candidate's experience supports them.
- Remove weak filler phrases such as "responsible for", "worked on", "helped with", "familiar with".
- If the candidate genuinely lacks experience for a gap, do NOT pretend they have it.
  Instead, reframe their closest real experience using the JD's vocabulary.

Return ONLY valid JSON.

Format:
{{
  "tailored_resume": ""
}}

Resume tailoring rules:
- Generate a complete tailored resume based only on the candidate's real resume.
- The tailored_resume field should contain the full resume text with these
  sections only when supported by the original resume:
    Contact
    Professional Summary
    Technical Skills
    Professional Experience
    Projects
    Education
    Certifications
- Do NOT include a Certifications section unless the original resume
  explicitly contains real certifications.
- Do NOT write planned, in progress, recommended, future, or suggested certifications.
- The final resume should look recruiter-ready, modern, concise, and strong
  enough for direct application.

Candidate Resume:
{data.resume_text}

Job Description:
{data.job_description}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an honest ATS analyst, recruiter, and resume strategist."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    parsed_response = json.loads(
        response.choices[0].message.content
    )

    tailored_resume = parsed_response.get("tailored_resume", "")

    # ----------------------------------------------------------------
    # Strip any "planned certifications" boilerplate the model added.
    # ----------------------------------------------------------------
    blocked_certification_phrases = [
        "certifications in progress",
        "certification in progress",
        "certifications planned",
        "certification planned",
        "planned certifications",
        "relevant certifications in progress",
        "can be added when obtained",
        "details can be added when obtained",
    ]

    if any(phrase in tailored_resume.lower() for phrase in blocked_certification_phrases):
        lines = tailored_resume.split("\n")
        cleaned_lines = []
        skip_certifications = False

        for line in lines:
            clean_line = line.strip()

            if clean_line.upper() == "CERTIFICATIONS":
                skip_certifications = True
                continue

            if skip_certifications and clean_line.isupper() and clean_line:
                skip_certifications = False

            if not skip_certifications:
                cleaned_lines.append(line)

        tailored_resume = "\n".join(cleaned_lines).strip()
        parsed_response["tailored_resume"] = tailored_resume

    # ----------------------------------------------------------------
    # STEP 3: Re-score the TAILORED resume with the EXACT same function.
    # This guarantees the projected score matches what the user would
    # see if they re-upload the new resume.
    # ----------------------------------------------------------------
    rescored = compare_resume_jd(
        ResumeJDRequest(
            resume_text=tailored_resume,
            job_description=data.job_description,
        )
    )
    new_score = rescored.get("match_score", 0)

    # ----------------------------------------------------------------
    # STEP 4: Safety net — never return a WORSE resume than the user uploaded.
    # If our tailoring lowered the score, return the original instead with
    # an honest explanation.
    # ----------------------------------------------------------------
    if new_score < original_score:
        parsed_response["tailored_resume"] = data.resume_text
        parsed_response["projected_ats_score_after_tailoring"] = original_score
        parsed_response["ats_score_reasoning"] = [
            f"Your resume is already well-aligned with this role at {original_score}/100.",
            "Our tailoring couldn't add meaningful value without inventing experience.",
            "Focus on adding real depth in these areas to push higher:",
            *[g for g in original_gaps[:3] if isinstance(g, str)],
        ]
        return parsed_response

    # ----------------------------------------------------------------
    # STEP 5: Normal path — return the tailored resume with its REAL score
    # and reasoning derived from the rescore (NOT fabricated by the rewrite).
    # ----------------------------------------------------------------
    reasoning_bullets = []

    verdict = rescored.get("recruiter_verdict", "").strip()
    if verdict:
        reasoning_bullets.append(verdict)

    for gap in rescored.get("critical_gaps", []):
        if isinstance(gap, str) and gap.strip():
            reasoning_bullets.append(gap.strip())

    parsed_response["projected_ats_score_after_tailoring"] = new_score
    parsed_response["ats_score_reasoning"] = reasoning_bullets

    return parsed_response

@app.post("/generate-cover-letter")
def generate_cover_letter(data: ResumeJDRequest):
    from datetime import datetime
    today_formatted = datetime.now().strftime("%B %d, %Y")

    prompt = f"""
You are an expert career writer and technical recruiter.

Your job is to write a tailored, professional cover letter for the candidate
based on their real resume and the target job description.

VERY IMPORTANT RULES:
- Do NOT invent fake experience, skills, companies, projects, or metrics.
- Only use what is genuinely supported by the candidate's resume.
- Match the tone to a modern, confident, professional job application.
- Keep it concise: 3 to 4 short paragraphs, under 350 words.
- Begin with a professional header containing the candidate's REAL contact
  information extracted from their resume: full name on the first line,
  then a single contact line with email, phone, and LinkedIn (whatever
  is actually present in the resume — omit anything missing, never invent).
- After the header, add this exact date on its own line: {today_formatted}
- Then write "Dear Hiring Manager," as the greeting (use the actual hiring
  manager's name only if it appears in the job description).
- Open the first paragraph with genuine interest in the specific role and
  company — never with "I am writing to apply" or similar clichés.
- Middle paragraphs connect the candidate's REAL experience to the JD's needs.
- Close with a confident call to action and a professional sign-off
  ("Sincerely," followed by the candidate's name).
- Avoid clichés like "team player", "hard worker", "go-getter".
- Avoid generic filler. Be specific to this candidate and this job.
- NEVER include placeholder text like [Company Name], [Your Name], [Date],
  or [Phone] — extract real data from the resume or omit the field entirely.
- Separate the header, date, greeting, body paragraphs, and sign-off with
  blank lines for readability.

Return ONLY valid JSON.

Format:

{{
  "cover_letter": ""
}}

The cover_letter field should contain the full cover letter as clean text,
with paragraphs separated by blank lines. No markdown, no bullet points.

Candidate Resume:
{data.resume_text}

Job Description:
{data.job_description}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an expert career writer and technical recruiter who writes honest, specific, compelling cover letters."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    parsed_response = json.loads(
        response.choices[0].message.content
    )

    return parsed_response

@app.post("/extract-jd-from-url")
async def extract_jd_from_url(data: dict):
    try:
        url = data.get("url")

        if not url:
            return {"error": "No URL provided"}

        headers = {
            "User-Agent": "Mozilla/5.0"
        }

        response = requests.get(url, headers=headers)

        soup = BeautifulSoup(response.text, "html.parser")

        text = soup.get_text(separator=" ", strip=True)

        cleaned_text = " ".join(text.split())

        return {
            "job_description": cleaned_text[:12000]
        }

    except Exception as e:
        return {
            "error": str(e)
        }
    
@app.post("/download-tailored-resume-docx")
def download_tailored_resume_docx(data: dict):
    tailored_resume = data.get("tailored_resume", "")

    file_id = str(uuid.uuid4())
    file_path = f"tailored_resume_{file_id}.docx"

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = doc.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    section_titles = [
        "PROFESSIONAL SUMMARY",
        "SUMMARY",
        "CORE COMPETENCIES",
        "SKILLS",
        "TECHNICAL SKILLS",
        "PROFESSIONAL EXPERIENCE",
        "WORK EXPERIENCE",
        "EXPERIENCE",
        "PROJECTS",
        "EDUCATION",
        "CERTIFICATIONS",
    ]

    lines = tailored_resume.split("\n")

    for index, line in enumerate(lines):
        clean_line = line.strip()

        if not clean_line:
            continue

        if index == 0:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.space_after = Pt(2)

            run = p.add_run(clean_line.upper())
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(17, 24, 39)
            continue

        if index == 1 and ("@" in clean_line or "|" in clean_line):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.space_after = Pt(8)

            run = p.add_run(clean_line)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(75, 85, 99)

            add_horizontal_line(p)

            continue

        if clean_line.upper() in section_titles:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)

            run = p.add_run(clean_line.upper())
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(37, 99, 235)

            add_horizontal_line(p)
            continue

        if clean_line.startswith("-") or clean_line.startswith("•"):
          p = doc.add_paragraph(style="List Bullet")
          p.paragraph_format.left_indent = Inches(0.28)
          p.paragraph_format.space_after = Pt(3.5)

          run = p.add_run(clean_line.lstrip("-• ").strip())
          run.font.size = Pt(10)
          run.font.color.rgb = RGBColor(31, 41, 55)
          continue

        is_role_line = (
            " | " in clean_line
            or " — " in clean_line
            or " - " in clean_line and any(char.isdigit() for char in clean_line)
        )

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)

        run = p.add_run(clean_line)
        run.font.size = Pt(10.5)

        if is_role_line:
         run.bold = True
         run.font.size = Pt(11)
         run.font.color.rgb = RGBColor(17, 24, 39)
         p.paragraph_format.space_before = Pt(5)
         p.paragraph_format.space_after = Pt(2)
    else:
         run.font.color.rgb = RGBColor(31, 41, 55)

    doc.save(file_path)

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="tailored_resume.docx"
    )

@app.post("/download-cover-letter-docx")
def download_cover_letter_docx(data: dict):
    cover_letter = data.get("cover_letter", "")

    file_id = str(uuid.uuid4())
    file_path = f"cover_letter_{file_id}.docx"

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    # Pre-process: extract name (first line) and contact line (second non-empty
    # line if it looks like contact info — contains @ or | or phone-ish text).
    # The AI is instructed to produce: NAME / CONTACT / DATE / DEAR... / BODY / SIGN-OFF.
    raw_lines = [line.strip() for line in cover_letter.split("\n")]
    non_empty = [l for l in raw_lines if l]

    name_line = non_empty[0] if non_empty else ""

    contact_line = ""
    if len(non_empty) > 1:
        candidate = non_empty[1]
        if "@" in candidate or "|" in candidate or "·" in candidate:
            contact_line = candidate

    # Normalize the contact line: convert pipes to middle dots for a softer look.
    if contact_line:
        contact_line = " · ".join(part.strip() for part in contact_line.replace("·", "|").split("|"))

    # Detect sign-off block so we can bold the closing name.
    # Convention: a line equal to "Sincerely," (or similar) followed by a name line.
    signoff_words = {"sincerely,", "best regards,", "regards,", "best,", "thanks,"}

    rendered_header = False
    skip_next_if_matches = None

    for index, line in enumerate(raw_lines):
        clean_line = line.strip()

        # Render the styled header once, replacing the first two content lines
        if not rendered_header and clean_line == name_line:
            # Name: large, bold, dark
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(name_line)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(17, 24, 39)

            # Contact: smaller, muted gray
            if contact_line:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(contact_line)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(107, 114, 128)

                # Horizontal divider
                add_horizontal_line(p)

            rendered_header = True
            # If the contact line was the second non-empty line, we'll need to skip it
            # when we encounter it in the regular loop.
            if contact_line and len(non_empty) > 1:
                skip_next_if_matches = non_empty[1]
            continue

        # Skip the contact line on its actual loop iteration (we already rendered it)
        if skip_next_if_matches and clean_line == skip_next_if_matches:
            skip_next_if_matches = None
            continue

        if not clean_line:
            continue

        # Detect "Sincerely," etc. and bold the next line (the closing name)
        if clean_line.lower() in signoff_words:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(clean_line)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(75, 85, 99)
            continue

        # Closing name: if the previous non-empty line was a sign-off,
        # render this line bold.
        prev_non_empty = next(
            (l for l in reversed(raw_lines[:index]) if l.strip()), ""
        )
        if prev_non_empty.lower() in signoff_words:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(17, 24, 39)
            continue

        # Regular body paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.4
        run = p.add_run(clean_line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(55, 65, 81)

    doc.save(file_path)

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="cover_letter.docx"
    )

@app.post("/download-tailored-resume-pdf")
def download_tailored_resume_pdf(data: dict):
    tailored_resume = data.get("tailored_resume", "")

    file_id = str(uuid.uuid4())
    file_path = f"tailored_resume_{file_id}.pdf"

    doc = SimpleDocTemplate(
        file_path,
        pagesize=letter,
        rightMargin=42,
        leftMargin=42,
        topMargin=34,
        bottomMargin=34,
    )

    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "NameStyle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=19,
        leading=23,
        alignment=TA_CENTER,
        textColor="#111827",
        spaceAfter=4,
    )

    contact_style = ParagraphStyle(
        "ContactStyle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.8,
        leading=11,
        alignment=TA_CENTER,
        textColor="#4B5563",
        spaceAfter=10,
    )

    heading_style = ParagraphStyle(
        "HeadingStyle",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10.8,
        leading=13,
        textColor="#2563EB",
        spaceBefore=9,
        spaceAfter=4,
    )

    normal_style = ParagraphStyle(
        "NormalResumeStyle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.2,
        leading=11.8,
        textColor="#1F2937",
        spaceAfter=4.5,
    )

    role_style = ParagraphStyle(
        "RoleStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12.5,
        textColor="#111827",
        spaceBefore=6,
        spaceAfter=3,
    )

    bullet_style = ParagraphStyle(
        "BulletResumeStyle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.1,
        leading=11.5,
        textColor="#1F2937",
        leftIndent=13,
        firstLineIndent=-7,
        spaceAfter=4.5,
    )

    story = []

    section_titles = [
        "PROFESSIONAL SUMMARY",
        "SUMMARY",
        "CORE COMPETENCIES",
        "SKILLS",
        "TECHNICAL SKILLS",
        "PROFESSIONAL EXPERIENCE",
        "WORK EXPERIENCE",
        "EXPERIENCE",
        "PROJECTS",
        "EDUCATION",
        "CERTIFICATIONS",
    ]

    lines = tailored_resume.split("\n")

    for index, line in enumerate(lines):
        clean_line = line.strip()

        if not clean_line:
            story.append(Spacer(1, 4))
            continue

        safe_line = (
            clean_line
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

        if index == 0:
            story.append(Paragraph(f"<b>{safe_line.upper()}</b>", name_style))
            continue

        if index == 1 and ("@" in clean_line or "|" in clean_line):
            story.append(Paragraph(safe_line, contact_style))
            continue

        if clean_line.upper() in section_titles:
            story.append(Paragraph(f"<b>{safe_line.upper()}</b>", heading_style))

            story.append(
                HRFlowable(
                    width="100%",
                    thickness=1,
                    color="#D1D5DB",
                    spaceBefore=2,
                    spaceAfter=6,
                )
            )

            continue

        if clean_line.startswith("-") or clean_line.startswith("•"):
            bullet_text = safe_line.lstrip("-• ").strip()
            story.append(Paragraph(f"• {bullet_text}", bullet_style))
            continue

        is_role_line = (
            " | " in clean_line
            or " — " in clean_line
            or (" - " in clean_line and any(char.isdigit() for char in clean_line))
        )

        if is_role_line:
            story.append(Paragraph(safe_line, role_style))
        else:
            story.append(Paragraph(safe_line, normal_style))


    doc.build(story)

    return FileResponse(
        file_path,
        media_type="application/pdf",
        filename="tailored_resume.pdf"
    )

@app.post("/download-cover-letter-pdf")
def download_cover_letter_pdf(data: dict):
    cover_letter = data.get("cover_letter", "")

    file_id = str(uuid.uuid4())
    file_path = f"cover_letter_{file_id}.pdf"

    doc = SimpleDocTemplate(
        file_path,
        pagesize=letter,
        rightMargin=64,
        leftMargin=64,
        topMargin=58,
        bottomMargin=58,
    )

    styles = getSampleStyleSheet()

    # Name header: bold, dark, large
    name_style = ParagraphStyle(
        "CoverLetterName",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor="#111827",
        spaceAfter=2,
    )

    # Contact line: small, muted
    contact_style = ParagraphStyle(
        "CoverLetterContact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        textColor="#6B7280",
        spaceAfter=6,
    )

    # Body paragraphs: readable, slightly muted
    letter_style = ParagraphStyle(
        "CoverLetterBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=17,
        textColor="#374151",
        spaceAfter=11,
    )

    # Sign-off line itself ("Sincerely,")
    signoff_style = ParagraphStyle(
        "CoverLetterSignoff",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=15,
        textColor="#4B5563",
        spaceBefore=8,
        spaceAfter=2,
    )

    # Closing name: bold, dark — mirrors the header name
    closing_name_style = ParagraphStyle(
        "CoverLetterClosingName",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=15,
        textColor="#111827",
        spaceAfter=10,
    )

    story = []

    # Pre-process: extract name and contact line (same logic as DOCX version)
    raw_lines = [line.strip() for line in cover_letter.split("\n")]
    non_empty = [l for l in raw_lines if l]

    name_line = non_empty[0] if non_empty else ""

    contact_line = ""
    if len(non_empty) > 1:
        candidate = non_empty[1]
        if "@" in candidate or "|" in candidate or "·" in candidate:
            contact_line = candidate

    # Normalize pipes to middle dots
    if contact_line:
        contact_line = " · ".join(part.strip() for part in contact_line.replace("·", "|").split("|"))

    signoff_words = {"sincerely,", "best regards,", "regards,", "best,", "thanks,"}

    rendered_header = False
    skip_next_if_matches = None

    def escape(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    for index, line in enumerate(raw_lines):
        clean_line = line.strip()

        # Render the styled header once
        if not rendered_header and clean_line == name_line:
            story.append(Paragraph(escape(name_line), name_style))
            if contact_line:
                story.append(Paragraph(escape(contact_line), contact_style))
                # Horizontal divider line
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=0.5,
                        color="#D1D5DB",
                        spaceBefore=2,
                        spaceAfter=10,
                    )
                )
            rendered_header = True
            if contact_line and len(non_empty) > 1:
                skip_next_if_matches = non_empty[1]
            continue

        # Skip the contact line on its real iteration (already rendered above)
        if skip_next_if_matches and clean_line == skip_next_if_matches:
            skip_next_if_matches = None
            continue

        if not clean_line:
            story.append(Spacer(1, 4))
            continue

        # Sign-off line
        if clean_line.lower() in signoff_words:
            story.append(Paragraph(escape(clean_line), signoff_style))
            continue

        # Closing name (line after a sign-off)
        prev_non_empty = next(
            (l for l in reversed(raw_lines[:index]) if l.strip()), ""
        )
        if prev_non_empty.lower() in signoff_words:
            story.append(Paragraph(escape(clean_line), closing_name_style))
            continue

        # Regular body paragraph
        story.append(Paragraph(escape(clean_line), letter_style))

    doc.build(story)

    return FileResponse(
        file_path,
        media_type="application/pdf",
        filename="cover_letter.pdf"
    )